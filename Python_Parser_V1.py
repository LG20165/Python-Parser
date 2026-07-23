from docx import Document
import os
import json
from collections import defaultdict


# ↓↓↓↓ HEADINGS AND FOOTERS ↓↓↓↓
#---------------- DETECTING HEADING LEVEL ----------------
def heading_level(p):
    if p.style is None:
        return 0
    
    name = p.style.name.lower()

    if "heading 1" in name:
        return 1
    if "heading 2" in name:
        return 2
    if "heading 3" in name:
        return 3
    return 0

#---------------- HEADER/FOOTER EXTRACTION ----------------
def parse_headers_and_footers(doc):
    blocks = []

    for i, section in enumerate(doc.sections):
        header = section.header
        footer = section.footer

        for p in header.paragraphs:
            if p.text.strip():
                blocks.append({
                    "type": "header",
                    "text": p.text.strip(),
                    "section": i
                })

        for p in footer.paragraphs:
            if p.text.strip():
                blocks.append({
                    "type": "footer",
                    "text": p.text.strip(),
                    "section": i
                })

    return blocks



# ↓↓↓↓ PARAGRAPH TEXT ↓↓↓↓
# ---------------- MAIN BODY PARSER ----------------
def parse_body(doc):
    blocks = []
    numbering_state= {}

    for p in doc.paragraphs:
        text = p.text.strip()
        if paragraph_contains_image(p):
            blocks.append({
                "type": "image",
                "image_file": None,
                "caption": None,
                "width": None,
                "height": None,
                "position": None,
                "page": None
            })
            continue

        if not text:
            continue

        kind = classify_paragraph(p)
        if kind == "heading":
            blocks.append({
                "type": "heading",
                "text": text,
                "level": heading_level(p)
            })

        elif kind == "list_item":
            number, level = get_numbering(p, numbering_state)
            blocks.append({
                "type": "list_item",
                "text": text,
                "numbering": number,
                "level": level
            })

        else:
            blocks.append({
                "type": "paragraph",
                "text": text,
                "style": p.style.name if p.style else None
            })
    return blocks



# ↓↓↓↓ LISTS ↓↓↓↓
#---------------- LIST DETECTION----------------
def classify_paragraph(p):
    pPr = p._p.pPr

    if pPr is not None and pPr.numPr is not None:
        return "list_item"
    
    if p.style and "Heading" in p.style.name:
        return "heading"
    return "paragraph"

#---------------- NUMBERING EXTRACTION----------------
def get_numbering(p, state):
    pPr = p._p.pPr

    if pPr is None or pPr.numPr is None:
        return None, None

    numId = pPr.numPr.numId.val
    ilvl = pPr.numPr.ilvl.val

    key = (numId, ilvl)
    if key not in state:
        state[key] = 1
    else:
        state[key] += 1
    return f"{state[key]}.", ilvl

#---------------- DETECT NUMBERED PARAGRAPHS ----------------
def is_list_paragraph(p):
    pPr = p._p.pPr
    return pPr is not None and pPr.numPr is not None



# ↓↓↓↓ IMAGING ↓↓↓↓
# ---------------- EXTRACT IMAGE SIZE ----------------
def get_image_size(drawing):
    extent = drawing[0].xpath('.//wp:extent', namespaces={
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    })
    
    if extent:
        cx = int[0].get("cx")
        cy = int[0].get("cy")
        return {
            "width": cx,
            "height": cy
        }
    
    return None

# ---------------- IMAGE EXTRACTION ----------------
def extract_images_and_metadata(doc):

    images = []

    for rel in doc.part._rels.values():
        if "image" in rel.target_ref:
            # Addds all of the metadata about the image to the dictionary of the image
            images.append({
                "type": "image",
                "image_file": os.path.basename(rel.target_ref),
                "file_path": rel.target_ref,
                "rel_id": rel.rId,
                "caption": None,
                "width": None,
                "height": None,
                "position": None,
                "page": None
             })

    return images

# ---------------- IMAGE DETECTION ----------------

def paragraph_contains_image(paragraph):
    for run in paragraph.runs:
        if "graphic" in run._element.xml:
            return True
    return False

#---------------- ENRICHMENT STAGE ----------------
def enrich_blocks(blocks, images):
    image_index = 0
    for block in blocks:
        # Uses XML code to see if the block is an image (stored as "figure"), and if so,
        # it searches if there are any more images left in our list of images created in the previous function
        # then in the dictionary it sets the "image_file" key to the actual filename of the image,
        # and then says that its done another image by increasing the index by 1
        if block["type"] == "figure" and image_index < len(images):
            block["image_file"] = images[image_index]["filename"]
            image_index += 1
        # if no page number was clarified it just assumes page 1
        if "page" in block and block["page"] is None:
            block["page"] = 1
    return blocks



# ↓↓↓↓ CAPTIONS ↓↓↓↓
# ---------------- CAPTION DETECTION ----------------
def is_caption(paragraph):
    style = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip().lower()

    if "Caption" in style:
        return True

    if text.startswith("figure") or text.startswith("fig.") or text.startswith("image") or text.startswith("Abbildung"):
        return True

    return False

# ---------------- CAPTION ATTACHING---------
def attach_captions(blocks):
    for i, block in enumerate(blocks):
        if block["type"] != "image":
            continue

        if i + 1 >= len(blocks):
            continue

        next_block = blocks[i + 1]
        if next_block["type"] != "paragraph":
            continue

        class FakeParagraph:
            def __init__(self, text):
                self.text = text
                self.style = None
        
        if is_caption(FakeParagraph(next_block["text"])):
            block["caption"] = next_block["text"]



# ---------------- TABLE PARSING ----------------
def parse_tables(doc):
    tables_out = []
    for table in doc.tables:
        grid = []
        for row in table.rows:
            grid.append([cell.text.strip() for cell in row.cells])
        tables_out.append({
            "type": "table",
            "rows": len(table.rows),
            "cols": len(table.rows[0].cells) if table.rows else 0,
            "data": grid,
            "page": None
        })
    return tables_out


# ↓↓↓↓ PUTTING IT ALL TOGETHER ↓↓↓↓
# ---------------- MAIN BODY PARSER ----------------
def parse_docx(docx_path):
    doc = Document(docx_path)
    numbering_state = defaultdict(lambda: defaultdict(int))
    blocks = []

    blocks += parse_body(doc)
    blocks += parse_tables(doc)
    blocks += extract_images_and_metadata(doc)
    blocks += parse_headers_and_footers(doc)

    attach_captions(blocks)

    return {"blocks": blocks}


# ---------------- RUN ----------------

if __name__ == "__main__":
    path = r"C:\Users\Hogan\Downloads\PROD-WI-07-DE_Copper Blocks vorbereiten_old.docx"

    result = parse_docx(path)

    with open(r"C:\Users\Hogan\OneDrive - Neoscan Solutions GmbH\Parser Code\parsed_output_PROD-WI-07-DE_Copper Blocks vorbereiten_old.json", "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    print("Done → parsed_output_dsl.json")