from docx import Document
import os
import json
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
import re

#---------------- LIST HANDLING STATE ----------------
# Track numbering state globally across the document body loop
list_state = {}

#---------------- ORDERING PARAGRAPHS AND TABLES ----------------
def iter_document_blocks(doc, image_lookup):
    index = 0
    body = doc.element.body

    image_counter_state = {"count": 0}

    def has_image(element):
        return len(element.xpath('.//w:drawing')or element.xpath('.//w:pict')) > 0

    for child in body.iterchildren():
        if child.tag.endswith('p'):
            para = Paragraph(child, doc)
            is_img = has_image(child)
            if not para.text.strip() and not is_img:
                continue
            
            if is_img:
                actual_image_id = image_counter_state["count"]
                image_counter_state["count"] += 1

                # find the relationship ID in the XML
                blips = child.xpath(".//a:blip")
                filename = None
                target = None

                if blips:
                    rid = blips[0].get("{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed")
                    # finds real file info using image_lookup
                    if rid and rid in image_lookup:
                        filename = image_lookup[rid]["filename"]
                        target = image_lookup[rid]["target"]
                meta = {
                    "style": para.style.name if para.style else None,
                    "layout": {},
                    "image_id": actual_image_id,
                    "filename": filename,
                    "target": target
                }
                yield make_block(index, "image", text=para.text.strip(), meta=meta)
                index += 1
                continue
            # initialize meta as a variable early so it can be added to later
            meta = {
                "style": para.style.name if para.style else None,
                "layout": {}
            }
            # Determine explicit paragraph type (heading, list_item, or paragraph)
            derived_type = classify_paragraph(para)
            style_name = para.style.name.lower() if para.style and para.style.name else ""
            if "bullet" in style_name:
                derived_type = "bullet_list_item"
            elif "list" in style_name:
                derived_type = "list_item"
            # Extract numbering based upon the item type (regular paragraph or header)
            num_prefix, ilvl = None, None
            # starts us with text stripped of all whitespace
            display_text = para.text.strip()

            if derived_type == "list_item":
                num_prefix, ilvl = get_numbering(para, list_state)
                if num_prefix:
                    meta["list_prefix"] = num_prefix
                    meta["indent_level"] = ilvl
                    display_text = f"{num_prefix} {display_text}"
                else:
                    meta["indent_level"] = 0
            elif derived_type == "heading":
                lvl = heading_level(para)
                meta["heading_level"] = lvl
                # either gets manual numbering or calculates it
                num_prefix = get_heading_prefix(lvl, heading_state)
                if num_prefix:
                    meta["list_prefix"] = num_prefix
                    display_text = f"{num_prefix} {display_text}"
            # processses inline run formatting without losing text modifications
            filtered_runs, global_formatting = analyze_paragraph_formatting(para)
            

            if global_formatting:
                meta["formatting"] = global_formatting
            
            if filtered_runs:
                meta["runs"] = filtered_runs

            yield make_block(index, derived_type, text=display_text, meta=meta)
            index += 1

        elif child.tag.endswith('tbl'):
            table = Table(child, doc)
            image_blocks = process_image_table(
                table,
                image_lookup,
                index,
                image_counter_state
            )

            if image_blocks:
                for img in image_blocks:
                    yield img
                index += len(image_blocks)
                continue

            grid = []
            for row in table.rows:
                grid.append([cell.text.strip() for cell in row.cells])

            yield make_block(
                index,
                "table",
                text="",
                meta = {
                    "rows": len(table.rows),
                    "cols": len(table.rows[0].cells) if table.rows else 0,
                    "raw_matrix": grid
                }
            )
            index += 1

#---------------- DETECTING STYLE AND FORMATTING CHANGES ----------------
def get_run_style(run):
    style = {}
    if run.bold:
        style["bold"] = True
    if run.italic:
        style["italic"] = True
    if run.underline:
        style["underline"] = True
    if run.font.highlight_color is not None:
        style["highlight"] = getattr(run.font.highlight_color, 'name', str(run.font.highlight_color))
    return style

#---------------- ANALYZE PARAGRAPH RUNS ---------------- 
def analyze_paragraph_formatting(paragraph):
    valid_runs = [r for r in paragraph.runs if r.text]
    if not valid_runs:
        return None, None

    run_styles = [get_run_style(r) for r in valid_runs]
    is_completely_plain = all(not style for style in run_styles)
    if is_completely_plain:
        return None, None
    is_uniform = all(style == run_styles[0] for style in run_styles)
    global_formatting = {"bold": False, "italic": False, "underline": False, "highlight": False}
    filtered_runs = []

    for run, style in zip(valid_runs, run_styles):
        if style.get("bold"): global_formatting["bold"] = True
        if style.get("italic"): global_formatting["italic"] = True
        if style.get("underline"): global_formatting["underline"] = True
        if style.get("highlight"): global_formatting["highlight"] = True

        if style:
            filtered_runs.append({"text": run.text, **style})
    
    global_formatting = {k: v for k, v in global_formatting.items() if v}
    if is_uniform:
        return None, global_formatting
    
    return filtered_runs, global_formatting


#---------------- MAKING BLOCKS ----------------
def make_block(index, block_type, text=None, meta=None):
    return {
        "index": index,
        "type": block_type,
        "text": text,
        "meta": meta or {},
    }


#---------------- GETTING HEADERS AND FOOTERS ----------------
def extract_headers_footers(doc):
    blocks = []
    seen_parts = set()
    for section in doc.sections:
        hf_pairs = [
            (section.header, "dynamic_header", "default"),
            (section.first_page_header, "dynamic_header", "first_page"),
            (section.even_page_header, "dynamic_header", "even_page"),
            (section.footer, "dynamic_footer", "default"),
            (section.first_page_footer, "dynamic_footer", "first_page"),
            (section.even_page_footer, "dynamic_footer", "even_page")
        ]

        for hf_object, hf_type, sub_type in hf_pairs:
            if not hf_object:
                continue

            part_id = hf_object.part.partname
            if part_id in seen_parts:
                continue
            has_content = False
            hf_combined_text = []

            if hf_object.paragraphs:
                p_text = "\n".join([p.text for p in hf_object.paragraphs if p.text.strip()])
                if p_text:
                    hf_combined_text.append(p_text)
                    has_content = True

            if hf_object.tables:
                for table in hf_object.tables:
                    table_rows = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_rows.append(" | ".join(row_text))
                    if table_rows:
                        hf_combined_text.append("\n".join(table_rows))
                        has_content = True

            if has_content:
                blocks.append({
                    "index": None,
                    "type": hf_type,
                    "text": "\n".join(hf_combined_text),
                    "meta": {"style": None, "subtype": sub_type},
                })
                seen_parts.add(part_id)
    return blocks

#---------------- EXTRACTING INFO FROM HEADER ----------------
def extract_header_metadata(text):
    # parses the entire document and looks for information needded in the .TXT file
    # takes out all whitespace and random formatting and just lookes for the name right after our key phrases
    draft_trigger_phrases = ["draft v", "draft: ","version", "version:"]
    draft_phrase = None
    text_to_check = text.lower()
    for phrase in draft_trigger_phrases:
        if phrase in text_to_check:
            draft_phrase = phrase
            break

    metadata = {
        "title": None,
        "subtitle": None,
        "wi_number": None,
        "doc_id": None,
        "draft": None,
        "language": None 
    }

    # \s* matches spaces and newlines (\n), ([^|]+) uses the "|" character in the JSON file as its bound to search
    patterns = {
    # Accounts for regular hyphens (-), en-dashes (–), or em-dashes (—)
    "wi_number": r"prod[-–—]wi[-–—]\s*([a-z0-9]+)", 
    "doc_id": r"docns[-–—]\s*([0-9-]+)",
    "subtitle": r"confidential\s*[\s|]\s*([^\n|]+)",
    }

    # is only going to add somethng to the draft if it finds one of the key phrases
    if draft_phrase:
        patterns["draft"] = rf"{re.escape(draft_phrase)}\s*(.*)"

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            # .strip() cleans up lingering spaces
            metadata[key] = match.group(1).strip()

    # seeing if a work instruction number was found
    if metadata["wi_number"]:
        metadata["title"] = f"PROD-WI-{metadata['wi_number']}"

    if "arbeitsanweisung" in text.lower():
        metadata["language"] = "DE"
    elif "work instruction" in text.lower():
        metadata["language"] = "EN"
    else:
        metadata["language"] = "Unknown"
    return metadata

#---------------- EXTRACTING INFO FROM FOOTER ----------------
def extract_footer_metadata(text):
    # parses the entire document and looks for information needded in the .TXT file
    # takes out all whitespace and random formatting and just lookes for the name right after our key phrases
    metadata = {
        "author": None,
        "editor": None,
        "approver": None
    }

    # \s* matches spaces and newlines (\n), ([^|]+) uses the "|" character in the JSON file as its bound to search
    patterns = {
        # footer information
        "author": r"Erstellt von:\s*([^|]+)",
        "editor": r"Geprüft von:\s*([^|]+)",
        "approver": r"Freigegeben von:\s*([^|]+)"
    }

    for key, pattern in patterns.items():
        match = re.search(pattern, text, re.IGNORECASE)
        if match:
            # .strip() cleans up lingering spaces
            metadata[key] = match.group(1).strip()
    return metadata
#---------------- DETECTING HEADING LEVEL ----------------
def heading_level(p):
    if p.style is None or p.style.name is None:
        # double check if XML is missing heading level or style
        pPr = p._p.pPr
        if pPr is not None and pPr.outlineLvl is not None:
            return int(pPr.outlineLvl.val) + 1
        return 0
    name = p.style.name.lower()
    if "heading 1" in name or "überschrift 1" in name: return 1
    if "heading 2" in name or "überschrift 2" in name: return 2
    if "heading 3" in name or "überschrift 3" in name: return 3
    if "bullet" in name or p.style.name.startswith("List Bullet"):
        derived_type = "bullet_list_item"
    elif "list" in name or p.style.name.startswith("List Number"):
        derived_type = "numbered_list_item"

    # double checkign outline level tags
    pPr = p._p.pPr
    if pPr is not None and pPr.outlineLvl is not None:
        return int(pPr.outlineLvl.val) + 1
    return 0

#---------------- LISTING HEADING LEVELS ----------------
list_state = {}
heading_state = {1: 0, 2: 0, 3: 0}
def get_heading_prefix(level, state):
    if level <= 0 or level > 3:
        return None
    # increment the current level
    state[level] += 1

    # resets all of the sub-levels when the upper header moves on
    for l in range(level + 1, 4):
        state[l] = 0
    
    # constructs the prefix string
    parts = []
    for l in range(1, level + 1):
        parts.append(str(state[l])) 

    return ".".join(parts) + "."

#---------------- LIST DETECTION ----------------
def classify_paragraph(p):
    pPr = p._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return "list_item"
    text_lower = p.text.strip().lower()
    if text_lower.startswith("step") or text_lower.startswith("schritt"):
        return "list_item"
    if heading_level(p) > 0:
        return "heading"
    return "paragraph"

#---------------- NUMBERING EXTRACTION ----------------
def get_numbering(p, state):
    pPr = p._p.pPr
    if pPr is None or pPr.numPr is None:
        return None, None

    # Safely extract XML node values for numbering definitions
    numId = pPr.numPr.numId.val if hasattr(pPr.numPr.numId, 'val') else 0
    ilvl = pPr.numPr.ilvl.val if hasattr(pPr.numPr.ilvl, 'val') else 0

    key = (numId, ilvl)
    if key not in state:
        state[key] = 1
    else:
        state[key] += 1
    return f"{state[key]}.", ilvl

#---------------- IMAGING & CAPTIONS ----------------
def extract_images(doc):

    image_lookup = {}
    
    for rel_id, rel in doc.part.rels.items():
        if rel.reltype != RT.IMAGE:
            continue
        
        image_lookup[rel_id] = {
            "filename": rel.target_ref.split("/")[-1],
            "target": rel.target_ref
        }

    return image_lookup

#---------------- CAPTION PROBABILITY AND PAIRING ----------------
def caption_score(paragraph):
    if not paragraph["text"]:
        return -1
    score = 0
    text = paragraph["text"].strip()
    lower = text.lower()
    style = paragraph["meta"].get("style")
    if style == "Caption":
        score += 100
    if lower.startswith(("figure", "fig.", "abb", "abbildung")):
        score += 50
    if len(text) < 150:
        score += 10
    if len(text.split()) < 20:
        score += 5
    return score

#---------------- GETTING IMAGES FROM TABLE ----------------
def process_image_table(table, image_lookup, start_index, image_counter_state):
    image_blocks = []
    current_index = start_index

    images = []
    captions = []
    for row_idx, row in enumerate(table.rows):
        for col_idx, cell in enumerate(row.cells):
            for para in cell.paragraphs:
                p_xml = para._element
                blips = [node for node in p_xml.iter() if node.tag.endswith('blip')]
                if blips:
                    for blip in blips:
                        rid = blip.get(
                            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
                        )

                        if rid:
                            images.append({
                                "rid": rid,
                                "row": row_idx,
                                "col": col_idx
                            })
                text = para.text.strip()
                if not text:
                    continue

                field_text = ""
                fld_fields = [node for node in p_xml.iter() if node.tag.endswith('fldSimple')]
                if fld_fields:
                    digits = [node.text for node in fld_fields[0].iter() if node.tag.endswith('t') and node.text]
                    if digits:
                        field_text = "".join(digits)
                
                num_text = ""
                lvl_texts = [node for node in p_xml.iter() if node.tag.endswith('lvlText')]
                if lvl_texts:
                    num_text = lvl_texts[0].get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val', '')

                prefix = ""
                if field_text:
                    prefix = field_text
                elif num_text:
                    prefix = num_text.replace('%1', '1.').replace('%2', 'a.').strip()
                
                if prefix and not text.startswith(prefix):
                    if prefix.isdigit() and ("figure" in text.lower() or "abb" in text.lower()):
                        text = re.sub(r'^(Figure|Abb\.)', rf'\1 {prefix}', text, flags=re.IGNORECASE)
                    else:
                        text = f"{prefix} {text}"


                candidate = {
                    "text": text,
                    "meta": {
                        "style": para.style.name if para.style else None
                    }
                }

                score = caption_score(candidate)
                if score > 0:
                    captions.append({
                        "text": text,
                        "score": score,
                        "row": row_idx,
                        "col": col_idx
                })
    for image in images:
        best_caption = ""
        best_score = -999
        for cap in captions:
            row_distance = abs(image["row"] - cap["row"])
            col_distance = abs(image["col"] - cap["col"])
            distance = row_distance + col_distance
            adjusted_score = cap["score"] - (distance * 15)
            if adjusted_score > best_score:
                best_score = adjusted_score
                best_caption = cap["text"]
        if image["rid"] not in image_lookup:
            continue

        metadata = image_lookup[image["rid"]]
        
        actual_image_id = image_counter_state["count"]
        image_counter_state["count"] += 1

        image_blocks.append({
            "index": current_index,
            "type": "image",
            "meta": {
                "image_id": actual_image_id,
                "filename": metadata["filename"],
                "target": metadata["target"],
                "caption": best_caption
            }
        })
        current_index += 1
    return image_blocks

#---------------- COMBINING LONELY IMAGES AND CAPTIONS ----------------
def link_standalone_images_and_captions(blocks):
    # finding if unpaired images and captions near eachother can be paired
    for i in range(len(blocks)-1):
        current_block = blocks[i]
        next_block = blocks[i+1]

        # Case 1: if caption directly follows image
        if current_block.get("type") == "image" and next_block.get("type") == "caption":
            # actually links the two
            if "meta" not in current_block:
                current_block["meta"] = {}
            current_block["meta"]["caption_text"] = next_block.get("text", "").strip()
            next_block["is_paired"] = True
        
        # Case 2: if caption comes before the image but less likely so we do the case 1 first
        elif current_block.get("type") == "caption" and next_block.get("type") == "image":
            if "meta" not in next_block:
                next_block["meta"] = {}
            next_block["meta"]["caption_text"] = current_block.get("text", "").strip()
            current_block["is_paired"] = True
    return blocks

#---------------- XML DOC PARSER ----------------
def parse_docx(docx_path):
    doc = Document(docx_path)
    blocks = []
    image_lookup = extract_images(doc)
    # 1. Process Structural Layout
    blocks.extend(extract_headers_footers(doc))
    blocks.extend(iter_document_blocks(doc, image_lookup))

    total_blocks = len(blocks)

    # 2. Sequential Normalization
    for i, block in enumerate(blocks):
        block["index"] = i
        # ensures block["meta"] exists so we dont run into a KeyError
        if "meta" not in block or block["meta"] is None:
            block["meta"] = {}

        # prints the meta data from the header here
        if block["type"] == "dynamic_header" and block["text"]:
            header_meta = extract_header_metadata(block["text"])
            block["meta"].update(header_meta)

        # prints the meta data from the foooter here
        if block["type"] == "dynamic_footer" and block["text"]:
            footer_meta = extract_footer_metadata(block["text"])
            block["meta"].update(footer_meta)
    final_cleaned_blocks = link_standalone_images_and_captions(blocks)
    return {
        "blocks": final_cleaned_blocks
    }

# This system is put in place to test with single document parsing
# and is ONLY meant for testing
# you do need to change the file path for input and output
# to comment or uncomment this entire text, higlight the text then press Ctrl + /
# if __name__ == "__main__":

#     #need to change this exact individual file name to change the document your parsing
#     exact_filename = "PROD-WI-07-DE_Copper Blocks vorbereiten_old"
#     input_folder_path = r"C:\Users\Hogan\OneDrive - Neoscan Solutions GmbH\DocsToParse"
#     input_path = os.path.join(input_folder_path, f"{exact_filename}.docx")
#     result = parse_docx(input_path)

#     header_title = result["blocks"][0]["meta"].get("title", "Unknown_Title")

#     # ensures that the title will be able to always be saved to file explorer, even if
#     # forbidden characters appear in the header title
#     forbidden_chars = [":", "/", "\\", "|", "*", "?", '"', "<", ">"]
#     for char in forbidden_chars:
#         header_title = header_title.replace(char, "-")
    
#     # removes any excess spaces in the title
#     final_json_doc_name = header_title.strip()
#     # personalize this variable for your own personal file that you want to output to
#     output_folder_path = r"C:\Users\Hogan\OneDrive - Neoscan Solutions GmbH\ParsedOutputs"

#     out_path = os.path.join(output_folder_path, f"parsed_{final_json_doc_name}.json")
    
#     with open(out_path, "w", encoding="utf-8") as f:
#         json.dump(result, f, indent=2, ensure_ascii=False)

#     print("Done → Succesfully Paresed the Document")
#     print(f"Output saved to: {out_path}")

# this code is meant for the final parsing of all of the documents
# it parses an entire folder of documents and outputs them to a seperate folder
# also still need to enter the file path for input and output
# have to use # and not """ since \ will still get processed, therefore breaking out any comments
# at least besides the pound sign
# to comment or uncomment this entire text, higlight the text then press Ctrl + /

if __name__ == "__main__":
    # this is the input folder you will be using where all of the .docx files are sitting
    input_folder = r"C:\Users\Hogan\OneDrive - Neoscan Solutions GmbH\DocsToParse"

    # this is the output file which starts empty and is where all of the .json files will end up
    output_folder = r"C:\Users\Hogan\OneDrive - Neoscan Solutions GmbH\ParsedOutputs"

    # is checking to make sure the folder exists, prevents everything from crashing
    os.makedirs(output_folder, exist_ok=True)

    # finds .docx files
    all_files = [f for f in os.listdir(input_folder) if f.endswith(".docx") and not f.startswith("~$")]
    print(f"Found {len(all_files)} documents to parse. Starting batch process ...\n")

    for file_name in all_files:
        full_input_path = os.path.join(input_folder, file_name)
        list_state.clear()
        heading_state.clear()
        heading_state.update({1: 0, 2: 0, 3: 0})
        print(f"Parsing: {file_name} --> ", end="", flush=True)

        try:
            #atually runs the parser
            result = parse_docx(full_input_path)
            # extracting title using header metadata first, then fallback to the file base name
            blocks = result.get("blocks", [])
            header_title = None
            for block in blocks:
                if block.get("type") == "dynamic_header" and isinstance(block.get("meta"), dict):
                    extracted_title = block["meta"].get("title")
                    if extracted_title:
                        header_title = str(extracted_title)
                        break
            if not header_title:
                header_title = os.path.splitext(file_name)[0] or "Unknown_Title"

            # ensures that the title will be able to always be saved to file explorer, even if
            # forbidden characters appear in the header title
            forbidden_chars = [":", "/", "\\", "|", "*", "?", '"', "<", ">"]
            for char in forbidden_chars:
                header_title = header_title.replace(char, "-")            

            final_json_doc_name = header_title.strip()
            full_output_path = os.path.join(output_folder, f"parsed_{final_json_doc_name}.json")

            # now actually saves the JSON file using the correct name
            with open(full_output_path, "w", encoding="utf-8") as f:
               json.dump(result, f, indent=2, ensure_ascii=False)
            print("SUCCESS ✅")

        except Exception as e:
            print(f"FAILED ❌ (Error: {e})")
    print("\nBatch processing complete! All documents have been parsed.")