from docx import Document
import os
import json
from collections import defaultdict
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum.text import WD_COLOR_INDEX


# ↓↓↓↓ KEEPING ORIGINAL WORD ORDER ↓↓↓↓
#---------------- ORDERING PARAGRAPHS AND TABLES ----------------
def iter_document_blocks(doc):
    index = 0
    body = doc.element.body

    for child in body.iterchildren():

        if child.tag.endswith('p'):
            para = Paragraph(child,doc)
            yield make_block(
                index,
                "paragraph",
                text=para.text,
                meta={
                    "style": para.style.name if para.style else None,
                    "runs": parse_runs(para),
                    "formatting": get_text_formatting(para),
                    "layout": {},
                }
            )
            index += 1

        elif child.tag.endswith('tbl'):
            table = Table(child,doc)
            table_text = []
            for row in table.rows:
                table_text.append(
                    [cell.text for cell in row.cells]
                )

            yield make_block(
                index,
                "table",
                text = str(table_text),
                meta={"rows": len(table.rows)}
            )
            index += 1


#---------------- HIGHLIGHT, BOLD, ITALIC, UNDERLINE ----------------
def get_text_formatting(paragraph):
    formatting={
        "bold": False,
        "italic": False,
        "underline": False,
        "highlight": False
    }

    for run in paragraph.runs:
        if run.bold:
            formatting["bold"] = True
        if run.italic:
            formatting["italic"] = True
        if run.underline:
            formatting["underline"] = True
        if run.font.highlight_color:
            formatting["highlight"] = True
    return formatting


#---------------- CREATING RUNS ----------------
def parse_runs(paragraph):
    runs = []

    for run in paragraph.runs:
        highlight = None
        if run.font.highlight_color is not None:
            highlight = getattr(run.font.highlight_color, 'name', str(run.font.highlight_color))
        runs.append({
            "text": run.text,
            "bold": bool(run.bold),
            "italic": bool(run.italic),
            "underline": bool(run.underline),
            "highlight": highlight,
            "font": run.font.name,
            "size": (
                run.font.size.pt
                if run.font.size
                else None
            )
        })
    return runs

#---------------- MAKING BLOCKS ----------------
def make_block(index, block_type, text = None, meta=None):
    return {
        "id": index,
        "index": index,
        "type": block_type,
        "text": text,
        "meta": meta or {},
        "links": []
    }


# ---------------- TABLE PARSING ----------------
def parse_tables(doc):
    tables_out = []
    for table in doc.tables:
        grid = []
        for row in table.rows:
            grid.append([cell.text.strip() for cell in row.cells])
        tables_out.append({
            "type": "table",
            "rows": len(table.rows),
            "cols": len(table.rows[0].cells) if table.rows else 0,
            "data": grid,
        })
    return tables_out


#---------------- INFERING DOCUMENT LAYOUT----------------
def infer_layout_position(index, total):
    ratio = index/max(total, 1)

    if ratio < 0.2:
        return "top"
    elif ratio < 0.8:
        return "midddle"
    else:
        return "bottom"




# ↓↓↓↓ HEADINGS AND FOOTERS ↓↓↓↓
#---------------- GETTING REAL HEADERS AND FOOTERS ----------------
def extract_headers_footers(doc):
    blocks = []
    headers = {
        "first_page": [],
        "even_page": [],
        "odd_page": [],
        "default": []
    }

    footers = {
        "first_page": [],
        "even_page": [],
        "odd_page": [],
        "default": []
    }

    for section in doc.sections:
# finding if there is anything on the default page
        if section.header:
            headers["default"].append({
                "type": "header",
                "subtype": "default",
                "text": "\n".join([p.text for p in section.header.paragraphs if p.text.strip()])
            })
        if section.footer:
            footers["default"].append({
                "type": "footer",
                "subtype": "default",
                "text": "\n".join([p.text for p in section.footer.paragraphs if p.text.strip()])
            })

# finding if there is anything on the first page
        if section.header:
            headers["first_page"].append({
                "type": "header",
                "subtype": "first_page",
                "text": "\n".join([p.text for p in section.header.paragraphs if p.text.strip()])
            })
        if section.footer:
            footers["first_page"].append({
                "type": "footer",
                "subtype": "first_page",
                "text": "\n".join([p.text for p in section.footer.paragraphs if p.text.strip()])
            })

# finding if there is anything on the even pagess
        if section.header:
            headers["even_page"].append({
                "type": "header",
                "subtype": "even_page",
                "text": "\n".join([p.text for p in section.header.paragraphs if p.text.strip()])
            })
        if section.footer:
            footers["even_page"].append({
                "type": "footer",
                "subtype": "even_page",
                "text": "\n".join([p.text for p in section.footer.paragraphs if p.text.strip()])
            })

# finding if there is anything on the odd pages
        if section.header:
            headers["odd_page"].append({
                "type": "header",
                "subtype": "odd_page",
                "text": "\n".join([p.text for p in section.header.paragraphs if p.text.strip()])
            })
        if section.footer:
            footers["odd_page"].append({
                "type": "footer",
                "subtype": "odd_page",
                "text": "\n".join([p.text for p in section.footer.paragraphs if p.text.strip()])
            })

    return headers, footers


#---------------- DETECTING HEADING LEVEL ----------------
def heading_level(p):
    if p.style is None:
        return 0
    
    name = p.style.name.lower()

    if "heading 1" in name:
        return 1
    if "heading 2" in name:
        return 2
    if "heading 3" in name:
        return 3
    return 0




# ↓↓↓↓ LISTS ↓↓↓↓
#---------------- LIST DETECTION----------------
def classify_paragraph(p):
    pPr = p._p.pPr

    if pPr is not None and pPr.numPr is not None:
        return "list_item"
    
    if p.style and "Heading" in p.style.name:
        return "heading"
    return "paragraph"


#---------------- NUMBERING EXTRACTION----------------
def get_numbering(p, state):
    pPr = p._p.pPr

    if pPr is None or pPr.numPr is None:
        return None, None

    numId = pPr.numPr.numId.val
    ilvl = pPr.numPr.ilvl.val

    key = (numId, ilvl)
    if key not in state:
        state[key] = 1
    else:
        state[key] += 1
    return f"{state[key]}.", ilvl


#---------------- DETECT NUMBERED PARAGRAPHS ----------------
def is_list_paragraph(p):
    pPr = p._p.pPr
    return pPr is not None and pPr.numPr is not None




# ↓↓↓↓ IMAGING ↓↓↓↓
#---------------- EXTRACTING IMAGES ----------------
def extract_images(doc):
    images = []
    image_id = 0

    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            images.append({
                "id": image_id,
                "filename": rel.target_ref.split("/")[-1],
                "target": rel.target_ref,
                "anchor_block": None
            })
            image_id += 1
    return images


# ---------------- EXTRACT IMAGE SIZE ----------------
def get_image_size(drawing):
    extent = drawing[0].xpath('.//wp:extent', namespaces={
        'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
    })
    
    if extent:
        cx = int(extent[0].get("cx"))
        cy = int(extent[0].get("cy"))
        return {
            "width": cx,
            "height": cy
        }
    
    return None




# ↓↓↓↓ CAPTIONS ↓↓↓↓
#---------------- DETECTING CAPTIONS WITH XML----------------
def find_caption_candidates(blocks, index, window = 2):
    candidates = []

    for i in range(max(0, index - window), min(len(blocks), index + window + 1)):
        if blocks[i]["type"] == "paragraph":
            text = (blocks[i]["text"] or "").strip()

            # is making the assumption that short text after a picture is likely a caption
            if 0 < len(text) < 200:
                candidates.append(blocks[i])

    return candidates


#---------------- LINKING IMAGES WITH TEXT----------------
def link_images_to_captions(blocks, images):
    relations = []

    for i, block in enumerate(blocks):
        if block["type"] != "image":
            continue
            
        captions = find_caption_candidates(blocks, i)
        for cap in captions:
            relations.append({
                "type": "image_caption",
                "image_id": block["id"],
                "caption_block_id": cap["id"]
            })

            block ["links"].append(cap["id"])
            cap["links"].append(block["id"])

    return relations


# ---------------- CAPTION DETECTION ----------------
def is_caption(paragraph):
    style = paragraph.style.name if paragraph.style else ""
    text = paragraph.text.strip().lower()

    if "Caption" in style:
        return True

    if text.startswith("figure") or text.startswith("fig.") or text.startswith("image") or text.startswith("Abbildung"):
        return True

    return False


# ---------------- CAPTION ATTACHING---------
def attach_captions(blocks):
    for i, block in enumerate(blocks):
        if block["type"] != "image":
            continue

        if i + 1 >= len(blocks):
            continue

        next_block = blocks[i + 1]
        if next_block["type"] != "paragraph":
            continue

        class FakeParagraph:
            def __init__(self, text):
                self.text = text
                self.style = None
        
        if is_caption(FakeParagraph(next_block["text"])):
            block["caption"] = next_block["text"]


#----------------SCORING CAPTION PROBABILITY ----------------
def caption_score(block):
    score = 0
    text = block["text"].lower()

    if text.startswith("Figure"):
        score += 10
    if text.startswith("Fig."):
        score += 10
    if text.startswith("Abb."):
        score += 10
    if text.startswith("Abbildung"):
        score += 10
    if len(text) < 150:
        score += 2
    if block["meta"]["style"] == "Caption":
        score += 100
    return score


#----------------PICKING MOST PROBABLE ----------------
def find_best_caption(blocks, image_index):
    best = None
    best_score = -1

    for i in range(max(0, image_index-3),
        min(len(blocks), image_index+4)):

        block = blocks[i]
        if block["type"] != "paragraph":
            continue

        score = caption_score(block)
        if score > best_score:
            best_score = score
            best = block
    return best




# ↓↓↓↓ PUTTING IT ALL TOGETHER ↓↓↓↓
#---------------- XML DOC PARSER ----------------
def parse_docx(docx_path):
    doc = Document(docx_path)
    blocks = []

    # headers and footers come first becuase they are at the top of the document
    blocks.extend(extract_headers_footers(doc))

    #now it parses the body
    blocks.extend(iter_document_blocks(doc))

    total_blocks = len(blocks)

    for i, block in enumerate(blocks):
        block["meta"]["layout"] = {
            "document_order": i,
            "position": infer_layout_position(i, total_blocks)
        }

    # now it parses the images which normally come after the text in the body
    images = extract_images(doc)

    relations = link_images_to_captions(blocks, images)

    return {
        "blocks": blocks,
        "images": images,
        "relations": relations
    }


# ---------------- RUN ----------------
if __name__ == "__main__":
    path = r"C:\Users\Hogan\Downloads\PROD-WI-07-DE_Copper Blocks vorbereiten_old.docx"
    doc=Document(path)

    result = parse_docx(path)

    with open(r"C:\Users\Hogan\OneDrive - Neoscan Solutions GmbH\Parser Code\parsed_output_PROD-WI-07-DE_Copper Blocks vorbereiten_old.json", "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    print("Done → parsed_output_dsl.json")

        