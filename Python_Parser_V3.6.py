from docx import Document
import os
import json
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

#---------------- LIST HANDLING STATE ----------------
# Track numbering state globally across the document body loop
list_state = {}

#---------------- ORDERING PARAGRAPHS AND TABLES ----------------
def iter_document_blocks(doc):
    index = 0
    body = doc.element.body

    def has_image(element):
        return len(element.xpath('.//w:drawing')or element.xpath('.//w:pict')) > 0

    for child in body.iterchildren():
        if child.tag.endswith('p'):
            para = Paragraph(child, doc)
            is_img = has_image(child)
            if not para.text.strip() and not is_img:
                continue
            
            if is_img:
                meta = {
                    "style": para.style.name if para.style else None,
                    "runs": parse_runs(para),
                    "layout": {},
                    "links": []
                }
                yield make_block(index, "image", text=para.text.strip(), meta=meta)
                index += 1
                continue
            # Determine explicit block type (heading, list_item, or paragraph)
            derived_type = classify_paragraph(para)
            
            # Extract numbering if it's a list item
            num_prefix, ilvl = None, None
            if derived_type == "list_item":
                num_prefix, ilvl = get_numbering(para, list_state)

            meta = {
                "style": para.style.name if para.style else None,
                "runs": parse_runs(para),
                "formatting": get_text_formatting(para),
                "layout": {}
            }

            # Inject Heading Levels if applicable
            if derived_type == "heading":
                meta["heading_level"] = heading_level(para)

            # Inject List Prefixes if applicable
            if derived_type == "list_item":
                meta["list_prefix"] = num_prefix
                meta["indent_level"] = ilvl
                # Prepend prefix to text so DeepSeek reads it clearly
                display_text = f"{num_prefix} {para.text}" if num_prefix else para.text
            else:
                display_text = para.text

            yield make_block(index, derived_type, text=display_text, meta=meta)
            index += 1

        elif child.tag.endswith('tbl'):
            table = Table(child, doc)
            grid = []
            nested_elements = []
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text.strip())
                    for p_element in cell._tc.xpath('.//w:p'):
                        cell_para = Paragraph(p_element, doc)

                        if has_image(p_element):
                            nested_elements.append({
                                "index": index,
                                "type": "image",
                                "text": cell_para.text.strip()
                            })
                        elif cell_para.text.strip():
                            nested_elements.append({
                                "index": index,
                                "type": "paragraph",
                                "text": cell_para.text.strip(),
                                "meta": {"style": cell_para.style.name if cell_para.style else "TableText"}
                            })
                grid.append(row_text)

            yield make_block(
                index,
                "table",
                text=json.dumps(grid, ensure_ascii=False),
                meta={
                    "rows": len(table.rows),
                    "cols": len(table.rows[0].cells) if table.rows else 0,
                    "raw_matrix": grid,
                    "elements": nested_elements
                }
            )
            index += 1

#---------------- HIGHLIGHT, BOLD, ITALIC, UNDERLINE ----------------
def get_text_formatting(paragraph):
    formatting = {"bold": False, "italic": False, "underline": False, "highlight": False}
    for run in paragraph.runs:
        if run.bold: formatting["bold"] = True
        if run.italic: formatting["italic"] = True
        if run.underline: formatting["underline"] = True
        if run.font.highlight_color is not None: formatting["highlight"] = True
    return formatting

#---------------- CREATING RUNS ----------------
def parse_runs(paragraph):
    runs = []
    for run in paragraph.runs:
        if not run.text:
            continue

        run_data = {"text": run.text}
        if run.bold:
            run_data["bold"] = True
        if run.italic:
            run_data["italic"] = True
        if run.underline:
            run_data["underline"] = True
        
        if run.font.highlight_color is not None:
            run_data["highlight"] = getattr(run.font.highlight_color, 'name', str(run.font.highlight_color))
        
        runs.append(run_data)
    return runs

#---------------- MAKING BLOCKS ----------------
def make_block(index, block_type, text=None, meta=None):
    return {
        "index": index,
        "type": block_type,
        "text": text,
        "meta": meta or {},
        "links": []
    }

#---------------- INFERRING DOCUMENT LAYOUT ----------------
def infer_layout_position(index, total):
    ratio = index / max(total, 1)
    if ratio < 0.2:
        return "top"
    elif ratio < 0.8:
        return "middle"
    else:
        return "bottom"

#---------------- GETTING HEADERS AND FOOTERS ----------------
def extract_headers_footers(doc):
    blocks = []
    seen_parts = set()
    for section in doc.sections:
        hf_pairs = [
            (section.header, "dynamic_header", "default"),
            (section.first_page_header, "dynamic_header", "first_page"),
            (section.even_page_header, "dynamic_header", "even_page"),
            (section.footer, "dynamic_footer", "default"),
            (section.first_page_footer, "dynamic_footer", "first_page"),
            (section.even_page_footer, "dynamic_footer", "even_page")
        ]

        for hf_object, hf_type, sub_type in hf_pairs:
            if not hf_object:
                continue

            part_id = hf_object.part.partname
            if part_id in seen_parts:
                continue
            has_content = False
            hf_combined_text = []

            if hf_object.paragraphs:
                p_text = "\n".join([p.text for p in hf_object.paragraphs if p.text.strip()])
                if p_text:
                    hf_combined_text.append(p_text)
                    has_content = True

            if hf_object.tables:
                for table in hf_object.tables:
                    table_rows = []
                    for row in table.rows:
                        row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                        if row_text:
                            table_rows.append(" | ".join(row_text))
                    if table_rows:
                        hf_combined_text.append("\n".join(table_rows))
                        has_content = True

            if has_content:
                blocks.append({
                    "index": None,
                    "type": hf_type,
                    "text": "\n".join(hf_combined_text),
                    "meta": {"style": None, "subtype": sub_type},
                    "links": []
                })
                seen_parts.add(part_id)
    return blocks

#---------------- DETECTING HEADING LEVEL ----------------
def heading_level(p):
    if p.style is None:
        return 0
    name = p.style.name.lower()
    if "heading 1" in name: return 1
    if "heading 2" in name: return 2
    if "heading 3" in name: return 3
    return 0

#---------------- LIST DETECTION ----------------
def classify_paragraph(p):
    pPr = p._p.pPr
    if pPr is not None and pPr.numPr is not None:
        return "list_item"
    if p.style and "Heading" in p.style.name:
        return "heading"
    return "paragraph"

#---------------- NUMBERING EXTRACTION ----------------
def get_numbering(p, state):
    pPr = p._p.pPr
    if pPr is None or pPr.numPr is None:
        return None, None

    # Safely extract XML node values for numbering definitions
    numId = pPr.numPr.numId.val if hasattr(pPr.numPr.numId, 'val') else 0
    ilvl = pPr.numPr.ilvl.val if hasattr(pPr.numPr.ilvl, 'val') else 0

    key = (numId, ilvl)
    if key not in state:
        state[key] = 1
    else:
        state[key] += 1
    return f"{state[key]}.", ilvl

#---------------- IMAGING & CAPTIONS ----------------
def extract_images(doc):
    images = []
    image_id = 0
    for rel in doc.part.rels.values():
        if rel.reltype == RT.IMAGE:
            images.append({
                "image_id": image_id,
                "filename": rel.target_ref.split("/")[-1],
                "target": rel.target_ref,
                "anchor_block": None
            })
            image_id += 1
    return images

def caption_score(block):
    if block["type"] != "paragraph":
        return -1
    score = 0
    text = (block["text"] or "").lower()
    if any(text.startswith(prefix) for prefix in ["figure", "fig.", "abb.", "abbildung"]):
        score += 10
    if len(text) < 150:
        score += 2
    if block["meta"].get("style") == "Caption":
        score += 100
    return score

def link_image_to_captions_with_table(blocks, raw_images):
    global_image_idx = 0
    relations = []
    for i, block in enumerate(blocks):
        # Case 1 which is finding captions inside of a table which is how the documents are structured
        if block.get("type") == "table":
            table_image = None
            table_caption_block = None
            highest_caption_score = -1

            nested_elements = block.get("meta", {}).get("elements", [])
            # goes through all cells inside the table block
            for element in nested_elements:
                if element.get("type") == "image":
                    table_image = element
                elif element.get("type") == "paragraph":
                    score = caption_score(element)
                    if score > highest_caption_score:
                        highest_caption_score = score
                        table_caption_block = element
            if table_image and table_caption_block and highest_caption_score > 0:
                if global_image_idx < len(raw_images):
                    relations.append({
                        "type": "image_caption",
                        "image_id": global_image_idx,
                        "caption_block_id": table_caption_block["index"]
                    })
                    if "links" not in table_caption_block:
                        table_caption_block["links"] = []
                    table_caption_block["links"].append(table_image["index"])
                    table_image["anchor_block"] = table_caption_block["index"]
                    global_image_idx += 1
            continue

        # Case 2 which is less likely that captions are found outside of a table
        if block.get("type") == "image":

            img_id = global_image_idx
            best_cap = None
            best_score = -1

            # window size is jut how far the parser will look to find a caption, and were going to assume that the caption is normally the next paragraph
            window_size = 2
            start_idx = max(0, i - window_size)
            end_idx = min(len(blocks), i + window_size + 1)

            for j in range(start_idx, end_idx):
                candidate_block = blocks[j]

                # calculating the score for only the paragraphs within range
                score = caption_score(candidate_block)

                distance = abs(i-j)
                adjusted_score = score - distance
                if adjusted_score > best_score:
                    best_score = adjusted_score
                    best_cap = candidate_block
            
            # actually links the best caption in range to the picture
            if best_cap and best_score > 0:
                if img_id < len(raw_images):
                    relations.append({
                        "type": "image_caption",
                        "image_id": img_id,
                        "caption_block_id": best_cap["index"]
                    })
                
                    if "links" not in best_cap:
                        best_cap["links"] = []

                    best_cap["links"].append(img_id)
                    block["anchor_block"] = best_cap["index"]
    return relations

def generate_clean_image_list(all_blocks, raw_images):
    clean_images = []

    # extracts individual elements from the table
    flat_elements = []
    for block in all_blocks:
        if not isinstance(block, dict):
            continue
        if block.get("type") == "table":
            # gets the nested elements of the table out
            elements = block.get("meta", {}).get("elements", [])
            for el in elements:
                if isinstance(el, dict):
                    flat_elements.append(el)
        else:
            flat_elements.append(block)

    caption_lookup = {}
    for el in flat_elements:
        if not isinstance(el, dict):
            continue
        links = el.get("links", []) or el.get("meta", {}).get("links", [])
        if el.get("type") == "paragraph" and links:
            for linked_index in links:
                caption_lookup[linked_index] = el.get("text", "")

    layout_images = [el for el in flat_elements if isinstance(el, dict) and el.get("type") == "image"]

    for i, img in enumerate(raw_images):
        if not isinstance(img, dict):
            continue

        caption_text = ""
        block_index = img.get("anchor_block")

        if i < len(layout_images):
            matched_layout_node = layout_images[i]
            block_index = matched_layout_node.get("index")
            caption_text = caption_lookup.get(block_index, "")

        clean_images.append({
            "image_id": img.get("image_id"),
            "filename": img.get("filename"),
            "target": img.get("target"),
            "caption": caption_text,
            "index": block_index
        })

    return clean_images

#---------------- XML DOC PARSER ----------------
def parse_docx(docx_path):
    doc = Document(docx_path)
    blocks = []

    # 1. Process Structural Layout
    blocks.extend(extract_headers_footers(doc))
    blocks.extend(iter_document_blocks(doc))

    total_blocks = len(blocks)

    # 2. Sequential Normalization
    for i, block in enumerate(blocks):
        block["index"] = i
        block["meta"]["layout"] = {
            "document_order": i,
            "position": infer_layout_position(i, total_blocks)
        }
    
    raw_images = extract_images(doc)
    relations = link_image_to_captions_with_table(blocks, raw_images)

    final_cleaned_blocks = []
    table_image_ids = {rel["image_id"] for rel in relations}
    table_caption_ids = {rel["caption_block_id"] for rel in relations}

    for block in blocks:
        if block.get("type") == "table":
            nested_elements = block.get("meta", {}).get("elements", [])
            contains_linked_assets = any(
                el["index"] in table_image_ids or el["index"] in table_caption_ids
                for el in nested_elements
            )
            if contains_linked_assets:
                for el in nested_elements:
                    if el["type"] == "image":
                        final_cleaned_blocks.append(el)
                continue

        final_cleaned_blocks.append(block)
    cleaned_images = generate_clean_image_list(final_cleaned_blocks, raw_images)
    return {
        "blocks": final_cleaned_blocks,
        "images": cleaned_images,
        "relations": relations
    }

if __name__ == "__main__":
    path = r"C:\Users\Hogan\Downloads\PROD-WI-07-DE_Copper Blocks vorbereiten_old.docx"
    result = parse_docx(path)

    out_path = r"C:\Users\Hogan\OneDrive - Neoscan Solutions GmbH\Parser Code\parsed_output_PROD-WI-07-DE_Copper Blocks vorbereiten_old.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(result, f, indent=2, ensure_ascii=False)

    print("Done → parsed_output_dsl.json")